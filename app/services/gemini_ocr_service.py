"""Document parsing service using LiteLLM (OpenAI-compatible) API."""

import asyncio
import base64
import io
import json
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from dotenv import load_dotenv
from openai import AsyncOpenAI
from PIL import Image

try:
    import docx
except ImportError:  # pragma: no cover - environment dependent
    docx = None

try:
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover - environment dependent
    fitz = None

# Load environment variables
load_dotenv()

from app.core.config import get_settings
from app.core.exceptions import OCRProcessingError, ProcessingTimeoutError
from app.core.logging import app_logger
from app.models.schemas import (
    DocumentContent,
    DocumentTypeEnum,
    OCRResponse,
    ProcessingStatusEnum,
    QuestionChoice,
)
from app.services.litellm_params import build_litellm_extra_params


class GeminiOCRService:
    """Service for document parsing using LiteLLM (OpenAI-compatible) API."""

    def __init__(self):
        self.settings = get_settings()
        self.supported_formats = [
            "png",
            "jpg",
            "jpeg",
            "tiff",
            "bmp",
            "gif",
            "webp",
            "pdf",
            "docx",
            "doc",
        ]

        # Configure LiteLLM client (uses same model as chat)
        if not self.settings.litellm_api_key:
            raise OCRProcessingError("LITELLM_API_KEY is required for OCR processing")
        headers = {}
        if self.settings.litellm_site_url:
            headers["HTTP-Referer"] = self.settings.litellm_site_url
        if self.settings.litellm_site_name:
            headers["X-Title"] = self.settings.litellm_site_name
        self.async_client = AsyncOpenAI(
            api_key=self.settings.litellm_api_key,
            base_url=self.settings.litellm_base_url,
            default_headers=headers or None,
        )
        self.model = self.settings.litellm_model
        self.timeout = self.settings.gemini_timeout
        self.max_retries = max(1, self.settings.gemini_max_retries)

    @staticmethod
    def _normalize_space(text: str) -> str:
        return re.sub(r"\s+", " ", str(text or "")).strip()

    @classmethod
    def _normalize_block_text(cls, text: Any) -> str:
        value = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
        if not value.strip():
            return ""

        normalized_lines: List[str] = []
        previous_blank = False
        for line in value.split("\n"):
            cleaned = re.sub(r"\s+", " ", line).strip()
            if cleaned:
                normalized_lines.append(cleaned)
                previous_blank = False
            elif normalized_lines and not previous_blank:
                normalized_lines.append("")
                previous_blank = True

        return "\n".join(normalized_lines).strip()

    @classmethod
    def _strip_question_prefix(cls, text: str) -> str:
        value = cls._normalize_space(text)
        if not value:
            return ""
        # Remove leading question numbering such as "1. ", "ข้อ 2) ", "Question 3: "
        patterns = [
            r"^\s*ข้อ(?:ที่)?\s*\d+\s*[\.\):\-]?\s+",
            r"^\s*question\s*\d+\s*[\.\):\-]?\s+",
            r"^\s*\d+\s*[\.\):\-]\s+",
        ]
        cleaned = value
        for pattern in patterns:
            cleaned = re.sub(pattern, "", cleaned, flags=re.IGNORECASE)
        return cls._normalize_space(cleaned)

    @classmethod
    def _strip_choice_prefix(cls, text: str) -> str:
        value = cls._normalize_space(text)
        if not value:
            return ""
        # Remove choice markers like "A) ", "ก. ", "1) " from the start only
        pattern = r"^\s*(?:[A-Da-d]|[กขคงจฉ]|[1-9][0-9]?)\s*[\.\):\-]\s+"
        cleaned = value
        # Some OCR outputs duplicated markers (e.g. "1. 1) ..."), so strip repeatedly.
        while True:
            next_value = re.sub(pattern, "", cleaned)
            if next_value == cleaned:
                break
            cleaned = next_value
        return cls._normalize_space(cleaned)

    @classmethod
    def _question_explicitly_needs_context(cls, question_text: str) -> bool:
        normalized = cls._normalize_space(question_text).lower()
        if not normalized:
            return False

        explicit_patterns = (
            r"จาก(?:ข้อความ|บทความ|บทอ่าน|เรื่อง|ตาราง|กราฟ|แผนภูมิ|แผนภาพ|บทสนทนา|ประกาศ|โฆษณา|จดหมาย|ข้อมูล)(?:ต่อไปนี้|ข้างต้น|ด้านบน|นี้)?",
            r"(?:พิจารณา|อ่าน)(?:ข้อความ|บทความ|ข้อมูล|ตาราง|กราฟ|แผนภาพ|ประกาศ|โฆษณา)",
            r"(?:ตาม|อ้างอิง)ข้อความ(?:ข้างต้น|ด้านบน|ต่อไปนี้)?",
            r"\baccording to\b",
            r"\bbased on\b",
            r"\bfrom the (?:passage|text|table|graph|chart|dialogue|advertisement)\b",
            r"\brefer to the (?:passage|text|table|graph|chart)\b",
        )
        return any(
            re.search(pattern, normalized, flags=re.IGNORECASE)
            for pattern in explicit_patterns
        )

    @classmethod
    def _looks_like_promotional_context(cls, context_text: str) -> bool:
        normalized = cls._normalize_space(context_text).lower()
        if not normalized:
            return False

        if re.search(r"https?://|www\.", normalized):
            return True

        signal_score = 0
        if re.search(
            r"(line\s*id|line[:\s]|@\w{2,}|facebook|instagram|ig[:\s])", normalized
        ):
            signal_score += 1
        if re.search(r"(tel[:\s]|โทร[:\s]?|ติดต่อ|contact)", normalized):
            signal_score += 1
        if re.search(r"(?:\+?\d[\d\-\s]{7,}\d)", normalized):
            signal_score += 1
        if re.search(r"(\.com|\.net|\.org|\.co\.th)", normalized):
            signal_score += 1
        return signal_score >= 2

    @classmethod
    def _has_question_context_token_overlap(
        cls,
        question_text: str,
        choices: List[str],
        shared_context: str,
    ) -> bool:
        question_blob = cls._normalize_space(
            " ".join(
                [question_text, *[str(choice or "") for choice in (choices or [])]]
            )
        )
        context_blob = cls._normalize_space(shared_context)
        if not question_blob or not context_blob:
            return False

        token_pattern = r"[A-Za-z0-9ก-๙]{2,}"
        question_tokens = set(re.findall(token_pattern, question_blob.lower()))
        context_tokens = set(re.findall(token_pattern, context_blob.lower()))
        if not question_tokens or not context_tokens:
            return False

        overlap = question_tokens & context_tokens
        overlap_ratio = len(overlap) / max(1, len(question_tokens))
        return len(overlap) >= 2 or overlap_ratio >= 0.18

    @classmethod
    def _resolve_question_context(
        cls, question: Dict[str, Any], content: Optional[Dict[str, Any]]
    ) -> str:
        if not isinstance(question, dict):
            return ""

        question_text = cls._normalize_block_text(
            question.get("question") or question.get("text") or question.get("prompt")
        )

        direct_candidates = [
            question.get("context"),
            question.get("question_context"),
            question.get("questionContext"),
            question.get("shared_context"),
            question.get("sharedContext"),
            question.get("passage"),
            question.get("reading_passage"),
            question.get("reference_text"),
            question.get("instructions"),
            question.get("instruction"),
        ]
        direct_context = next(
            (
                cls._normalize_block_text(candidate)
                for candidate in direct_candidates
                if cls._normalize_block_text(candidate)
            ),
            "",
        )

        section_title = cls._normalize_block_text(
            question.get("section_title")
            or question.get("sectionTitle")
            or question.get("section")
        )
        structured_fallback_context = ""
        description_fallback_context = ""
        if isinstance(content, dict):
            structured_candidates = [
                content.get("context"),
                content.get("question_context"),
                content.get("shared_context"),
                content.get("passage"),
                content.get("instructions"),
                content.get("instruction"),
            ]
            structured_fallback_context = next(
                (
                    cls._normalize_block_text(candidate)
                    for candidate in structured_candidates
                    if cls._normalize_block_text(candidate)
                ),
                "",
            )

            # `content.description` is often generic document text; use it only when
            # the question explicitly references an external passage/table/etc.
            if question_text and cls._question_explicitly_needs_context(question_text):
                description_fallback_context = cls._normalize_block_text(
                    content.get("description")
                )

        resolved = (
            direct_context
            or structured_fallback_context
            or description_fallback_context
        )
        if section_title and resolved and section_title not in resolved:
            resolved = f"{section_title}\n{resolved}".strip()

        if resolved and question_text and resolved == question_text:
            return ""
        return resolved

    def _resolve_context_classifier_model(self) -> str:
        configured_model = str(
            self.settings.chat_context_classifier_model or ""
        ).strip()
        if configured_model:
            return configured_model
        return self.model

    @staticmethod
    def _coerce_bool(value: Any) -> bool:
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return value != 0
        text = str(value or "").strip().lower()
        return text in {
            "true",
            "1",
            "yes",
            "y",
            "related",
            "need",
            "required",
            "include",
            "ใช่",
            "ต้อง",
        }

    @staticmethod
    def _coerce_confidence(value: Any) -> float:
        if value is None:
            return 0.0
        text = str(value).strip()
        if not text:
            return 0.0
        try:
            confidence = float(text)
        except (TypeError, ValueError):
            return 0.0
        if confidence > 1.0 and confidence <= 100.0:
            confidence = confidence / 100.0
        return max(0.0, min(1.0, confidence))

    @staticmethod
    def _safe_extract_json_object(raw: str) -> Optional[Dict[str, Any]]:
        text = (raw or "").strip()
        if not text:
            return None
        if text.startswith("```"):
            text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
            text = re.sub(r"\s*```$", "", text)
        try:
            parsed = json.loads(text)
            return parsed if isinstance(parsed, dict) else None
        except Exception:
            pass
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            try:
                parsed = json.loads(text[start : end + 1])
                return parsed if isinstance(parsed, dict) else None
            except Exception:
                return None
        return None

    async def _call_litellm_prompt(
        self,
        prompt: str,
        max_tokens: Optional[int] = None,
        model_name: Optional[str] = None,
        include_reasoning_effort: bool = True,
    ) -> str:
        """Call LiteLLM text model with a plain prompt."""
        resolved_model = str(model_name or "").strip() or self.model
        for attempt in range(self.max_retries):
            try:
                extra_params = (
                    build_litellm_extra_params(
                        resolved_model,
                        reasoning_effort=self.settings.litellm_reasoning_effort,
                    )
                    if include_reasoning_effort
                    else {}
                )
                if isinstance(max_tokens, int) and max_tokens > 0:
                    extra_params["max_tokens"] = max_tokens
                    extra_params["max_completion_tokens"] = max_tokens
                response = await self.async_client.chat.completions.create(
                    model=resolved_model,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0,
                    timeout=self.timeout,
                    **extra_params,
                )
                content = (
                    response.choices[0].message.content if response.choices else ""
                )
                if not content:
                    raise OCRProcessingError("Empty response from LiteLLM API")
                return str(content).strip()
            except Exception as e:
                if attempt >= self.max_retries - 1:
                    raise OCRProcessingError(f"LiteLLM API call failed: {e}")
                await asyncio.sleep(1.0)

    async def _llm_question_context_relevance_classifier(
        self,
        question_text: str,
        choices: List[str],
        shared_context: str,
    ) -> Tuple[bool, bool, str, float]:
        """Classify if context is needed and whether it aligns with question/choices."""
        compact_choices = [
            re.sub(r"\s+", " ", str(choice or "")).strip()
            for choice in (choices or [])
            if re.sub(r"\s+", " ", str(choice or "")).strip()
        ][:6]
        rendered_choices = "\n".join(
            f"{chr(65 + idx)}) {choice}" for idx, choice in enumerate(compact_choices)
        )
        prompt = (
            "คุณคือระบบตรวจสอบบริบทข้อสอบจาก OCR\n"
            "ให้ตอบ JSON เพียงบรรทัดเดียวเท่านั้น (ห้าม markdown/code fence)\n"
            'รูปแบบ: {"requires_context": true|false, "context_related": true|false, "confidence": 0.0-1.0, "reason": "short"}\n'
            "เกณฑ์:\n"
            "- requires_context=true เมื่อโจทย์/ช้อยส์ต้องใช้ข้อความบริบทร่วมเพื่อทำข้อสอบ\n"
            "- requires_context=false เมื่อโจทย์และช้อยส์มีข้อมูลพอครบเอง\n"
            "- context_related=true เมื่อบริบทที่ให้มาช่วยตอบโจทย์/ช้อยส์นี้ได้จริง\n"
            "- context_related=false เมื่อบริบทคนละเรื่อง ไม่ช่วยแก้โจทย์นี้ หรือขัดกับโจทย์/ช้อยส์\n"
            "- confidence เป็นตัวเลข 0 ถึง 1\n\n"
            f"โจทย์:\n{question_text[:1200]}\n\n"
            f"ช้อยส์:\n{rendered_choices[:1200]}\n\n"
            f"บริบทประกอบโจทย์:\n{shared_context[:1800]}"
        )
        configured_max_tokens = int(self.settings.chat_context_classifier_max_tokens)
        base_max_tokens = max(120, configured_max_tokens)
        retry_max_tokens = min(
            max(base_max_tokens * 2, base_max_tokens + 120),
            1200,
        )
        token_budgets = [base_max_tokens]
        if retry_max_tokens > base_max_tokens:
            token_budgets.append(retry_max_tokens)

        last_reason = "invalid_classifier_payload"
        try:
            for idx, token_budget in enumerate(token_budgets):
                raw = await self._call_litellm_prompt(
                    prompt,
                    max_tokens=token_budget,
                    model_name=self._resolve_context_classifier_model(),
                    include_reasoning_effort=False,
                )
                parsed = self._safe_extract_json_object(raw)
                if isinstance(parsed, dict):
                    requires_raw = parsed.get("requires_context")
                    if requires_raw is None:
                        requires_raw = parsed.get("need_context")
                    related_raw = parsed.get("context_related")
                    if related_raw is None:
                        related_raw = parsed.get("aligned")
                    requires_context = self._coerce_bool(requires_raw)
                    context_related = self._coerce_bool(related_raw)
                    confidence_raw = parsed.get("confidence")
                    if confidence_raw is None:
                        confidence_raw = parsed.get("score")
                    confidence = self._coerce_confidence(confidence_raw)
                    reason = str(parsed.get("reason") or "classified")
                    return requires_context, context_related, reason, confidence

                last_reason = "invalid_classifier_payload"
                if idx < len(token_budgets) - 1:
                    app_logger.warning(
                        "OCR context classifier returned invalid payload, retrying with larger token budget"
                    )
        except Exception as exc:
            app_logger.warning(f"OCR context relevance classifier failed: {exc}")
            last_reason = "context_classifier_failed"
        return True, True, last_reason, 0.0

    async def _should_keep_question_context(
        self,
        question_text: str,
        choices: List[str],
        shared_context: str,
    ) -> Tuple[bool, str]:
        """Decide if OCR context should be kept for a question."""
        if not str(shared_context or "").strip():
            return False, "empty_context"

        question_requires_context = self._question_explicitly_needs_context(
            question_text
        )
        has_token_overlap = self._has_question_context_token_overlap(
            question_text=question_text,
            choices=choices,
            shared_context=shared_context,
        )
        looks_promotional = self._looks_like_promotional_context(shared_context)

        if (
            looks_promotional
            and not question_requires_context
            and not has_token_overlap
        ):
            return False, "ocr_context_drop_promotional_or_contact"

        if not self.settings.chat_context_classifier_enabled:
            return True, "context_classifier_disabled_keep"

        (
            requires_context,
            context_related,
            reason,
            confidence,
        ) = await self._llm_question_context_relevance_classifier(
            question_text=question_text,
            choices=choices,
            shared_context=shared_context,
        )
        confidence_threshold = float(
            self.settings.chat_context_classifier_confidence_threshold
        )
        if confidence < confidence_threshold:
            if looks_promotional and not question_requires_context:
                return (
                    False,
                    (
                        "ocr_context_low_confidence_drop_promotional:"
                        f"{reason}:confidence={confidence:.2f}:threshold={confidence_threshold:.2f}"
                    ),
                )
            if question_requires_context or has_token_overlap:
                return (
                    True,
                    (
                        "ocr_context_low_confidence_keep_heuristic:"
                        f"{reason}:confidence={confidence:.2f}:threshold={confidence_threshold:.2f}"
                    ),
                )
            return (
                False,
                (
                    "ocr_context_low_confidence_drop:"
                    f"{reason}:confidence={confidence:.2f}:threshold={confidence_threshold:.2f}"
                ),
            )
        if not requires_context:
            return (
                False,
                f"ocr_context_not_required:{reason}:confidence={confidence:.2f}",
            )
        if not context_related:
            return (
                False,
                f"ocr_context_mismatch:{reason}:confidence={confidence:.2f}",
            )
        if (
            looks_promotional
            and not question_requires_context
            and not has_token_overlap
        ):
            return (
                False,
                f"ocr_context_drop_promotional_guard:{reason}:confidence={confidence:.2f}",
            )
        return True, f"ocr_context_keep:{reason}:confidence={confidence:.2f}"

    async def parse_document(
        self,
        file_path: Path,
        document_id: str,
        document_type: DocumentTypeEnum = DocumentTypeEnum.DOCUMENT,
        language: str = "auto",
        enhance_markdown: bool = True,
        extract_questions: bool = True,
        selected_pages: Optional[List[int]] = None,
    ) -> OCRResponse:
        """
        Parse document using LiteLLM (OpenAI-compatible) API.

        Args:
            file_path: Path to document file
            document_id: Unique document identifier
            document_type: Type of document being processed
            language: Primary language of the document
            enhance_markdown: Whether to enhance output with additional formatting
            extract_questions: Whether to extract questions and choices as JSON

        Returns:
            OCRResponse with parsing results

        Raises:
            OCRProcessingError: If parsing fails
        """
        start_time = datetime.utcnow()

        try:
            app_logger.info(
                f"Starting LiteLLM OCR processing for document {document_id}"
            )

            if not file_path.exists():
                raise OCRProcessingError(f"File not found: {file_path}")

            # Validate file format
            if not self.validate_file_format(file_path):
                raise OCRProcessingError(f"Unsupported file format: {file_path.suffix}")

            # Handle different file types
            file_extension = file_path.suffix.lower().lstrip(".")

            if file_extension in ["png", "jpg", "jpeg", "tiff", "bmp", "gif", "webp"]:
                # Handle image files
                image = Image.open(file_path)
                prompt = self._create_parsing_prompt(
                    document_type, language, extract_questions
                )

                app_logger.info(
                    f"Parsing image document: {file_path} with LiteLLM model {self.model}"
                )
                response = await self._call_litellm_vision(image, prompt)
            elif file_extension == "pdf":
                # Handle PDF files by converting pages to images and using Vision API
                prompt = self._create_parsing_prompt(
                    document_type, language, extract_questions
                )

                app_logger.info(
                    f"Parsing PDF document: {file_path} with LiteLLM Vision API"
                )
                response = await self._process_pdf_with_vision(
                    file_path, prompt, selected_pages
                )
            elif file_extension in ["docx", "doc"]:
                # Handle DOCX files
                text_content = await self._extract_text_from_docx(file_path)
                prompt = self._create_text_parsing_prompt(
                    document_type, language, extract_questions
                )

                app_logger.info(
                    f"Parsing DOCX document: {file_path} with LiteLLM model {self.model}"
                )
                response = await self._call_litellm_text(text_content, prompt)
            else:
                raise OCRProcessingError(f"Unsupported file format: {file_extension}")

            if not response or not response.strip():
                raise OCRProcessingError("Empty content extracted from document")

            # Parse the JSON response
            original_text = response
            structured_content = None
            questions_data = None

            # Parse the structured JSON response
            try:
                parsed_json = self._parse_json_response(response)
                if parsed_json:
                    # Debug: Log the parsed questions
                    app_logger.info(
                        f"Parsed questions from OCR: {parsed_json.get('questions', [])}"
                    )

                    # Validate and fix questions that might be missing choices
                    valid_questions = []
                    context_kept_count = 0
                    context_dropped_count = 0
                    context_screening_logs: List[str] = []
                    for q in parsed_json.get("questions", []):
                        if isinstance(q, dict) and "question" in q:
                            # Ensure choices exists and is a list
                            if "choices" not in q or not isinstance(q["choices"], list):
                                app_logger.warning(
                                    f"Question missing choices: {q.get('question', 'Unknown')}"
                                )
                                q["choices"] = []  # Add empty choices array

                            normalized_question = self._strip_question_prefix(
                                q.get("question", "")
                            )
                            normalized_context = self._resolve_question_context(
                                q,
                                parsed_json.get("content")
                                if isinstance(parsed_json.get("content"), dict)
                                else {},
                            )
                            normalized_choices = [
                                self._strip_choice_prefix(choice)
                                for choice in (q.get("choices") or [])
                            ]
                            normalized_choices = [
                                choice for choice in normalized_choices if choice
                            ]

                            if not normalized_question:
                                continue

                            normalized_question_payload = {
                                **q,
                                "question": normalized_question,
                                "choices": normalized_choices,
                            }
                            if normalized_context:
                                (
                                    keep_context,
                                    context_reason,
                                ) = await self._should_keep_question_context(
                                    question_text=normalized_question,
                                    choices=normalized_choices,
                                    shared_context=normalized_context,
                                )
                                if keep_context:
                                    normalized_question_payload[
                                        "context"
                                    ] = normalized_context
                                    context_kept_count += 1
                                else:
                                    context_dropped_count += 1
                                    context_screening_logs.append(
                                        f"drop:{context_reason}"
                                    )
                                    app_logger.info(
                                        "Dropping OCR question context (%s): question='%s' context='%s'",
                                        context_reason,
                                        normalized_question[:140],
                                        normalized_context[:180],
                                    )

                            valid_questions.append(normalized_question_payload)

                    structured_content = DocumentContent(
                        document_type=parsed_json.get(
                            "document_type", document_type.value
                        ),
                        title=parsed_json.get("title"),
                        content=parsed_json.get("content", {}),
                        questions=[QuestionChoice(**q) for q in valid_questions],
                        metadata={
                            **(
                                parsed_json.get("metadata")
                                if isinstance(parsed_json.get("metadata"), dict)
                                else {}
                            ),
                            "context_screening": {
                                "enabled": bool(
                                    self.settings.chat_context_classifier_enabled
                                ),
                                "kept_count": context_kept_count,
                                "dropped_count": context_dropped_count,
                                "sample_reasons": context_screening_logs[:20],
                            },
                        },
                    )

                    # Extract questions data for backward compatibility
                    if valid_questions:
                        questions_data = [QuestionChoice(**q) for q in valid_questions]

            except Exception as e:
                app_logger.error(f"Failed to parse JSON response: {e}")
                # Fallback to basic text processing
                structured_content = DocumentContent(
                    document_type=document_type.value,
                    title=None,
                    content={"raw_text": response, "sections": []},
                    questions=None,
                    metadata={"parsing_error": str(e)},
                )

            # Calculate processing time
            end_time = datetime.utcnow()
            processing_time_ms = int((end_time - start_time).total_seconds() * 1000)

            # Calculate confidence score based on content quality
            confidence_score = self._calculate_confidence_score(original_text)

            app_logger.info(
                f"LiteLLM OCR processing completed for document {document_id}"
            )

            return OCRResponse(
                document_id=document_id,
                status=ProcessingStatusEnum.COMPLETED,
                original_text=original_text,
                structured_content=structured_content,
                questions_data=questions_data,
                confidence_score=confidence_score,
                processing_time_ms=processing_time_ms,
                metadata={
                    "document_type": document_type.value,
                    "language": language,
                    "enhance_markdown": enhance_markdown,
                    "extract_questions": extract_questions,
                    "parser_type": "LiteLLM",
                    "model": self.model,
                    "file_size": file_path.stat().st_size,
                    "questions_count": len(questions_data) if questions_data else 0,
                    "storage_format": "json",
                },
            )

        except Exception as e:
            app_logger.error(
                f"LiteLLM OCR processing failed for document {document_id}: {e}"
            )

            # Calculate processing time even for errors
            end_time = datetime.utcnow()
            processing_time_ms = int((end_time - start_time).total_seconds() * 1000)

            return OCRResponse(
                document_id=document_id,
                status=ProcessingStatusEnum.ERROR,
                error_message=str(e),
                processing_time_ms=processing_time_ms,
                metadata={
                    "document_type": document_type.value,
                    "language": language,
                    "enhance_markdown": enhance_markdown,
                    "extract_questions": extract_questions,
                    "parser_type": "LiteLLM",
                    "model": self.model,
                    "error_type": type(e).__name__,
                    "storage_format": "json",
                },
            )

    def _create_parsing_prompt(
        self, document_type: DocumentTypeEnum, language: str, extract_questions: bool
    ) -> str:
        """Create appropriate prompt based on document type and requirements."""

        if document_type == DocumentTypeEnum.EXAM and extract_questions:
            return """Please analyze this exam document and extract all content in a structured JSON format.

Return the result as a JSON object with this structure:
{
  "document_type": "exam",
  "title": "document title if present",
  "content": {
    "description": "any introductory text or instructions",
    "sections": []
  },
  "questions": [
    {
      "context": "shared instruction, section heading, reading passage, table/chart description, or common stem needed to answer this question. Use empty string if none.",
      "question": "complete question text",
      "choices": ["choice 1", "choice 2", "choice 3", "choice 4"]
    }
  ]
}

Important requirements:
1. Put only the item-specific stem in "question"
2. Put all shared instructions, section labels, passages, table/chart descriptions, or common stems needed to answer in "context"
3. If multiple questions depend on the same shared text, repeat that text in each relevant question's "context"
4. MUST include all answer choices (A, B, C, D or ก, ข, ค, ง etc.) - EVERY question MUST have choices array
5. If a question has multiple choice answers, include ALL choices in the choices array
6. If no clear choices are visible, create an empty choices array []
7. Preserve the original language (Thai or English)
8. Return ONLY valid JSON format, no additional text
9. If no questions are found, return empty questions array
10. Include any overall exam instructions or descriptions in content.description
11. In "question", remove numbering prefixes like "1.", "2)", "ข้อ 3", "Question 4"
12. In each item of "choices", remove labels like "A.", "B)", "ก.", "1)", etc.
13. Do not drop reading passages or shared instructions for comprehension-style questions

CRITICAL: Every question object MUST have both "question" and "choices" fields. Include "context" whenever any supporting text is needed to answer."""

        elif document_type == DocumentTypeEnum.BOOK:
            return f"""Please extract all text from this book/textbook and return it as structured JSON.

Return the result as a JSON object with this structure:
{{
  "document_type": "book",
  "title": "book title if present",
  "content": {{
    "chapters": [
      {{
        "title": "chapter title",
        "sections": [
          {{
            "heading": "section heading",
            "content": "section text content",
            "subsections": []
          }}
        ]
      }}
    ],
    "raw_text": "full extracted text"
  }}
}}

Language: {language if language != 'auto' else 'auto-detect'}
Return ONLY valid JSON format, no additional text."""

        else:  # DOCUMENT
            return f"""Please extract all text from this document and return it as structured JSON.

Return the result as a JSON object with this structure:
{{
  "document_type": "document",
  "title": "document title if present",
  "content": {{
    "sections": [
      {{
        "heading": "section heading",
        "content": "section text content"
      }}
    ],
    "raw_text": "full extracted text"
  }}
}}

Language: {language if language != 'auto' else 'auto-detect'}
Return ONLY valid JSON format, no additional text."""

    def _create_text_parsing_prompt(
        self, document_type: DocumentTypeEnum, language: str, extract_questions: bool
    ) -> str:
        """Create appropriate prompt for text-based document parsing."""

        if document_type == DocumentTypeEnum.EXAM and extract_questions:
            return """Please analyze this exam document text and extract all content in a structured JSON format.

Return the result as a JSON object with this structure:
{
  "document_type": "exam",
  "title": "document title if present",
  "content": {
    "description": "any introductory text or instructions",
    "sections": []
  },
  "questions": [
    {
      "context": "shared instruction, section heading, reading passage, table/chart description, or common stem needed to answer this question. Use empty string if none.",
      "question": "complete question text",
      "choices": ["choice 1", "choice 2", "choice 3", "choice 4"]
    }
  ]
}

Important requirements:
1. Put only the item-specific stem in "question"
2. Put all shared instructions, section labels, passages, table/chart descriptions, or common stems needed to answer in "context"
3. If multiple questions depend on the same shared text, repeat that text in each relevant question's "context"
4. MUST include all answer choices (A, B, C, D or ก, ข, ค, ง etc.) - EVERY question MUST have choices array
5. If a question has multiple choice answers, include ALL choices in the choices array
6. If no clear choices are visible, create an empty choices array []
7. Preserve the original language (Thai or English)
8. Return ONLY valid JSON format, no additional text
9. If no questions are found, return empty questions array
10. Include any overall exam instructions or descriptions in content.description
11. In "question", remove numbering prefixes like "1.", "2)", "ข้อ 3", "Question 4"
12. In each item of "choices", remove labels like "A.", "B)", "ก.", "1)", etc.
13. Do not drop reading passages or shared instructions for comprehension-style questions

CRITICAL: Every question object MUST have both "question" and "choices" fields. Include "context" whenever any supporting text is needed to answer."""

        elif document_type == DocumentTypeEnum.BOOK:
            return f"""Please extract all text from this book/textbook document and return it as structured JSON.

Return the result as a JSON object with this structure:
{{
  "document_type": "book",
  "title": "book title if present",
  "content": {{
    "chapters": [
      {{
        "title": "chapter title",
        "sections": [
          {{
            "heading": "section heading",
            "content": "section text content",
            "subsections": []
          }}
        ]
      }}
    ],
    "raw_text": "full extracted text"
  }}
}}

Language: {language if language != 'auto' else 'auto-detect'}
Return ONLY valid JSON format, no additional text."""

        else:  # DOCUMENT
            return f"""Please extract all text from this document and return it as structured JSON.

Return the result as a JSON object with this structure:
{{
  "document_type": "document",
  "title": "document title if present",
  "content": {{
    "sections": [
      {{
        "heading": "section heading",
        "content": "section text content"
      }}
    ],
    "raw_text": "full extracted text"
  }}
}}

Language: {language if language != 'auto' else 'auto-detect'}
Return ONLY valid JSON format, no additional text."""

    async def _call_litellm_vision(self, image: Image.Image, prompt: str) -> str:
        """Call LiteLLM Vision API with the image and prompt."""
        buffered = io.BytesIO()
        image.save(buffered, format="PNG")
        image_b64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
        image_url = f"data:image/png;base64,{image_b64}"

        for attempt in range(self.max_retries):
            try:
                extra_params = build_litellm_extra_params(
                    self.model,
                    reasoning_effort=self.settings.litellm_reasoning_effort,
                )
                response = await self.async_client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": prompt},
                                {"type": "image_url", "image_url": {"url": image_url}},
                            ],
                        }
                    ],
                    temperature=0,
                    timeout=self.timeout,
                    **extra_params,
                )
                content = (
                    response.choices[0].message.content if response.choices else ""
                )
                if not content:
                    raise OCRProcessingError("Empty response from LiteLLM API")
                return content.strip()
            except Exception as e:
                if attempt >= self.max_retries - 1:
                    raise OCRProcessingError(f"LiteLLM API call failed: {e}")
                await asyncio.sleep(1.0)

    async def _call_litellm_text(self, text_content: str, prompt: str) -> str:
        """Call LiteLLM API with text content and prompt."""
        full_prompt = f"{prompt}\n\nDocument content:\n{text_content}"

        for attempt in range(self.max_retries):
            try:
                extra_params = build_litellm_extra_params(
                    self.model,
                    reasoning_effort=self.settings.litellm_reasoning_effort,
                )
                response = await self.async_client.chat.completions.create(
                    model=self.model,
                    messages=[{"role": "user", "content": full_prompt}],
                    temperature=0,
                    timeout=self.timeout,
                    **extra_params,
                )
                content = (
                    response.choices[0].message.content if response.choices else ""
                )
                if not content:
                    raise OCRProcessingError("Empty response from LiteLLM API")
                return content.strip()
            except Exception as e:
                if attempt >= self.max_retries - 1:
                    raise OCRProcessingError(f"LiteLLM API call failed: {e}")
                await asyncio.sleep(1.0)

    async def _extract_images_from_pdf(self, file_path: Path) -> List[Image.Image]:
        """Convert PDF pages to images for LiteLLM Vision API processing."""
        try:
            if fitz is None:
                raise OCRProcessingError("PyMuPDF (fitz) is not installed")

            def pdf_to_images():
                images = []

                # Use PyMuPDF to convert PDF pages to images
                doc = fitz.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)

                    # Convert page to image with high DPI for better OCR
                    matrix = fitz.Matrix(2.0, 2.0)  # 2x zoom for better resolution
                    pix = page.get_pixmap(matrix=matrix)

                    # Convert PyMuPDF pixmap to PIL Image
                    img_data = pix.tobytes("ppm")
                    img = Image.open(io.BytesIO(img_data))
                    images.append(img)

                doc.close()
                return images

            return await asyncio.to_thread(pdf_to_images)

        except Exception as e:
            app_logger.error(f"Failed to convert PDF to images {file_path}: {e}")
            raise OCRProcessingError(f"PDF to image conversion failed: {e}")

    async def _process_pdf_with_vision(
        self, file_path: Path, prompt: str, selected_pages: Optional[List[int]] = None
    ) -> str:
        """Process PDF by converting pages to images and using LiteLLM Vision API."""
        try:
            # Convert PDF pages to images
            images = await self._extract_images_from_pdf(file_path)

            if not images:
                raise OCRProcessingError("No images extracted from PDF")

            # Filter images to only selected pages if specified
            if selected_pages:
                # Convert to 0-based indexing and filter images
                selected_indices = [
                    p - 1 for p in selected_pages if 1 <= p <= len(images)
                ]
                filtered_images = [(i, images[i]) for i in selected_indices]
                app_logger.info(
                    f"Processing selected pages {selected_pages} (indices {selected_indices}) from PDF"
                )
            else:
                # Process all pages
                filtered_images = [(i, image) for i, image in enumerate(images)]
                app_logger.info(f"Processing all {len(images)} pages from PDF")

            if not filtered_images:
                raise OCRProcessingError("No valid PDF pages selected for OCR")

            try:
                configured_concurrency = int(self.settings.ocr_pdf_max_concurrency or 1)
            except (TypeError, ValueError):
                configured_concurrency = 1
            max_concurrency = max(1, configured_concurrency)
            concurrency = min(max_concurrency, len(filtered_images))
            semaphore = asyncio.Semaphore(concurrency)
            app_logger.info(
                f"Processing {len(filtered_images)} PDF page(s) with LiteLLM Vision "
                f"concurrency={concurrency}"
            )

            async def process_page(
                original_page_num: int, image: Image.Image
            ) -> Dict[str, Any]:
                display_page_num = original_page_num + 1
                app_logger.info(
                    f"Processing PDF page {display_page_num}/{len(images)} with LiteLLM Vision"
                )

                page_prompt = (
                    f"{prompt}\n\nThis is page {display_page_num} from a PDF document."
                )

                async with semaphore:
                    try:
                        response = await self._call_litellm_vision(image, page_prompt)
                    except Exception as page_error:
                        app_logger.warning(
                            f"Skipping PDF page {display_page_num}/{len(images)} "
                            f"after LiteLLM Vision failure: {page_error}"
                        )
                        return {
                            "page": display_page_num,
                            "response": None,
                            "error": str(page_error),
                        }

                    if response and response.strip():
                        return {
                            "page": display_page_num,
                            "response": response,
                            "error": None,
                        }

                    app_logger.warning(
                        f"Skipping PDF page {display_page_num}/{len(images)} "
                        "because LiteLLM Vision returned empty content"
                    )
                    return {
                        "page": display_page_num,
                        "response": None,
                        "error": "Empty response from LiteLLM API",
                    }

            page_results = await asyncio.gather(
                *[
                    process_page(original_page_num, image)
                    for original_page_num, image in filtered_images
                ]
            )
            page_results.sort(key=lambda item: item["page"])
            all_responses = [
                item["response"] for item in page_results if item.get("response")
            ]
            failed_pages = [
                {"page": item["page"], "error": item["error"]}
                for item in page_results
                if item.get("error")
            ]

            if failed_pages:
                app_logger.warning(
                    f"PDF vision skipped {len(failed_pages)} page(s): {failed_pages}"
                )

            if not all_responses:
                raise OCRProcessingError(
                    "No usable responses extracted from PDF pages"
                    + (f"; skipped pages: {failed_pages}" if failed_pages else "")
                )

            # Combine responses from all pages
            if len(all_responses) == 1:
                return all_responses[0]
            else:
                # For multiple pages, we need to merge the JSON responses
                return self._merge_multi_page_responses(all_responses)

        except Exception as e:
            app_logger.error(f"Failed to process PDF with vision API {file_path}: {e}")
            raise OCRProcessingError(f"PDF vision processing failed: {e}")

    def _merge_multi_page_responses(self, responses: List[str]) -> str:
        """Merge JSON responses from multiple PDF pages."""
        try:
            merged_questions = []
            merged_content = {"sections": [], "raw_text": ""}
            document_title = None
            document_type = "exam"

            for page_num, response in enumerate(responses):
                try:
                    page_data = self._parse_json_response(response)
                    if page_data:
                        # Collect title from first page
                        if page_num == 0 and page_data.get("title"):
                            document_title = page_data["title"]

                        # Merge questions
                        if page_data.get("questions"):
                            merged_questions.extend(page_data["questions"])

                        # Merge content
                        if page_data.get("content"):
                            content = page_data["content"]
                            if content.get("sections"):
                                merged_content["sections"].extend(content["sections"])
                            if content.get("raw_text"):
                                merged_content[
                                    "raw_text"
                                ] += f"\n--- Page {page_num + 1} ---\n{content['raw_text']}"
                            if content.get("description"):
                                if page_num == 0:
                                    merged_content["description"] = content[
                                        "description"
                                    ]

                except Exception as e:
                    app_logger.warning(
                        f"Failed to parse response from page {page_num + 1}: {e}"
                    )
                    # Include raw response as fallback
                    merged_content[
                        "raw_text"
                    ] += f"\n--- Page {page_num + 1} (Raw) ---\n{response}"

            # Create final merged response
            merged_response = {
                "document_type": document_type,
                "title": document_title or "Multi-page Document",
                "content": merged_content,
                "questions": merged_questions,
            }

            return json.dumps(merged_response, ensure_ascii=False, indent=2)

        except Exception as e:
            app_logger.error(f"Failed to merge multi-page responses: {e}")
            # Return concatenated raw responses as fallback
            return "\n\n--- PAGE BREAK ---\n\n".join(responses)

    async def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text content from DOCX file."""
        try:
            if docx is None:
                raise OCRProcessingError("python-docx is not installed")

            def extract_docx_text():
                doc = docx.Document(file_path)
                text_content = ""

                for paragraph in doc.paragraphs:
                    text_content += paragraph.text + "\n"

                # Extract text from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + " "
                        text_content += "\n"

                return text_content

            return await asyncio.to_thread(extract_docx_text)

        except Exception as e:
            app_logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            raise OCRProcessingError(f"DOCX text extraction failed: {e}")

    def _parse_json_response(self, response: str) -> Optional[Dict[str, Any]]:
        """Parse JSON response from Gemini API."""
        try:
            # Clean the response text
            response = response.strip()

            # Try to parse as JSON first
            if response.startswith("{") and response.endswith("}"):
                return json.loads(response)

            # Try to find JSON in the response
            import re

            # Look for JSON object pattern
            json_match = re.search(r"\{[\s\S]*\}", response, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                try:
                    return json.loads(json_str)
                except json.JSONDecodeError:
                    pass

            # Try to extract JSON from code blocks
            code_block_match = re.search(
                r"```(?:json)?\s*(\{[\s\S]*?\})\s*```", response, re.DOTALL
            )
            if code_block_match:
                json_str = code_block_match.group(1)
                try:
                    return json.loads(json_str)
                except json.JSONDecodeError:
                    pass

            # If no JSON found, return None
            app_logger.warning("No valid JSON found in Gemini response")
            return None

        except json.JSONDecodeError as e:
            app_logger.error(f"Failed to parse JSON from Gemini response: {e}")
            return None
        except Exception as e:
            app_logger.error(f"Failed to parse response: {e}")
            return None

    def _extract_questions_from_response(self, response: str) -> List[QuestionChoice]:
        """Extract questions and choices from Gemini response."""
        try:
            # Clean the response text
            response = response.strip()

            # Try to parse as JSON first
            if response.startswith("[") and response.endswith("]"):
                data = json.loads(response)
                return [QuestionChoice(**item) for item in data]

            # If not pure JSON, try to find JSON in the response
            import re

            # Look for JSON array pattern
            json_match = re.search(r"\[[\s\S]*?\]", response, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                try:
                    data = json.loads(json_str)
                    return [QuestionChoice(**item) for item in data]
                except json.JSONDecodeError:
                    pass

            # Try to extract JSON from code blocks
            code_block_match = re.search(
                r"```(?:json)?\s*(\[[\s\S]*?\])\s*```", response, re.DOTALL
            )
            if code_block_match:
                json_str = code_block_match.group(1)
                try:
                    data = json.loads(json_str)
                    return [QuestionChoice(**item) for item in data]
                except json.JSONDecodeError:
                    pass

            # If no JSON found, return empty list
            app_logger.warning(
                "No valid JSON found in Gemini response for question extraction"
            )
            return []

        except json.JSONDecodeError as e:
            app_logger.error(f"Failed to parse JSON from Gemini response: {e}")
            return []
        except Exception as e:
            app_logger.error(f"Failed to extract questions: {e}")
            return []

    def _generate_markdown_from_questions(
        self, questions_data: List[QuestionChoice]
    ) -> str:
        """Generate markdown content from extracted questions."""
        if not questions_data:
            return "# แบบทดสอบ\n\nไม่พบคำถามในเอกสาร"

        markdown_lines = ["# แบบทดสอบ\n"]

        for i, q_data in enumerate(questions_data, 1):
            if q_data.context:
                markdown_lines.append(f"{q_data.context}\n")
            markdown_lines.append(f"## ข้อ {i}: {q_data.question}\n")

            for j, choice in enumerate(q_data.choices):
                # Auto-detect language and use appropriate choice letters
                if any("\u0e00" <= c <= "\u0e7f" for c in q_data.question):
                    # Thai text detected, use Thai choice letters
                    choice_letter = chr(ord("ก") + j)
                else:
                    # Use English choice letters
                    choice_letter = chr(ord("A") + j)

                markdown_lines.append(f"{choice_letter}. {choice}")

            markdown_lines.append("")  # Add blank line between questions

        return "\n".join(markdown_lines)

    def _calculate_confidence_score(self, text: str) -> float:
        """Calculate confidence score based on extracted text quality."""
        if not text or not text.strip():
            return 0.0

        # Basic heuristics for confidence scoring
        score = 0.7  # Base score for Gemini (higher than OpenAI OCR)

        # Length factor
        text_length = len(text.strip())
        if text_length > 1000:
            score += 0.15
        elif text_length > 500:
            score += 0.1

        # Structure factor
        if any(marker in text for marker in ["#", "*", "-", "1.", "•", "|", "{", "["]):
            score += 0.1

        # Sentence structure factor
        sentences = text.split(".")
        if len(sentences) > 5:
            score += 0.05

        return min(0.95, score)

    def enhance_markdown_formatting(
        self, text: str, document_type: DocumentTypeEnum
    ) -> str:
        """Enhance markdown formatting based on document type."""
        if not text:
            return ""

        enhanced_text = text

        # Add document type header if not present
        if not text.strip().startswith("#"):
            type_names = {
                DocumentTypeEnum.DOCUMENT: "เอกสาร",
                DocumentTypeEnum.BOOK: "หนังสือ",
                DocumentTypeEnum.EXAM: "แบบทดสอบ",
            }
            header = f"# {type_names.get(document_type, 'เอกสาร')}\n\n"
            enhanced_text = header + text

        return enhanced_text

    async def get_parsing_status(self, job_id: str) -> Dict[str, Any]:
        """Get parsing status for a job (Gemini API is synchronous)."""
        return {
            "job_id": job_id,
            "status": "completed",
            "message": "Gemini API processes documents synchronously",
        }

    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return self.supported_formats.copy()

    def validate_file_format(self, file_path: Path) -> bool:
        """Validate if file format is supported."""
        file_extension = file_path.suffix.lower().lstrip(".")
        return file_extension in self.supported_formats
